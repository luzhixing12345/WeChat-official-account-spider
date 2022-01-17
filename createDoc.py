from docx import Document
from docx.shared import Inches

def make_doc(title,publish_time,doc):
    document = Document()

    #添加标题，并设置级别，范围：0 至 9，默认为1
    document.add_heading(title, 0)

    #添加段落，文本可以包含制表符（\t）、换行符（\n）或回车符（\r）等
    document.add_paragraph(publish_time)

    for i in range(len(doc)):
        if doc[i][1]=='' or doc[i][2]==0:
            continue
        if doc[i][0]=='subtitle':
            document.add_heading(doc[i][1], level=1)
        elif doc[i][0]=='content':
            document.add_paragraph(doc[i][1])
        elif doc[i][0]=='image':
            document.add_picture(doc[i][1],width=Inches(5.9))
        else:
            raise
    # document.add_heading('Heading, level 1', level=1)
    # document.add_paragraph('Intense quote', style='Intense Quote')

    # #添加项目列表（前面一个小圆点）
    # document.add_paragraph(
    #     'first item in unordered list', style='List Bullet'
    # )
    # document.add_paragraph('second item in unordered list', style='List Bullet')

    # #添加项目列表（前面数字）
    # document.add_paragraph('first item in ordered list', style='List Number')
    # document.add_paragraph('second item in ordered list', style='List Number')

    # #添加图片
    

    #保存.docx文档
    document.save('paper.docx')